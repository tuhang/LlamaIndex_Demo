"""
教案知识库管理模块
负责教案文档的加载、处理、向量化和检索
"""
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd
from docx import Document
import PyPDF2

from llama_index.core import (
    Document as LlamaDocument,
    VectorStoreIndex,
    StorageContext,
    Settings
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
import chromadb

from config import settings

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LessonPlanKnowledgeBase:
    """教案知识库管理类"""
    
    def __init__(self):
        """初始化知识库"""
        self.knowledge_base_dir = Path(settings.knowledge_base_dir)
        self.chroma_persist_dir = Path(settings.chroma_persist_dir)
        
        # 确保目录存在
        self.knowledge_base_dir.mkdir(parents=True, exist_ok=True)
        self.chroma_persist_dir.mkdir(parents=True, exist_ok=True)
        
        # 配置LlamaIndex
        Settings.embed_model = OpenAIEmbedding(
            model=settings.embedding_model,
            api_key=settings.openai_api_key,
            api_base=settings.openai_api_base
        )
        Settings.llm = OpenAI(
            model=settings.llm_model,
            api_key=settings.openai_api_key,
            api_base=settings.openai_api_base
        )
        
        # 初始化ChromaDB
        self.chroma_client = chromadb.PersistentClient(path=str(self.chroma_persist_dir))
        self.collection_name = "lesson_plans"
        
        # 初始化向量存储
        self._init_vector_store()
        
    def _init_vector_store(self):
        """初始化向量存储"""
        try:
            # 获取或创建collection
            self.chroma_collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name
            )
            
            # 创建ChromaVectorStore实例
            self.vector_store = ChromaVectorStore(chroma_collection=self.chroma_collection)
            
            # 创建存储上下文
            self.storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )
            
            logger.info("向量存储初始化成功")
            
        except Exception as e:
            logger.error(f"向量存储初始化失败: {e}")
            raise
    
    def load_documents(self, file_paths: List[str] = None) -> List[LlamaDocument]:
        """
        加载教案文档
        
        Args:
            file_paths: 指定文件路径列表，如果为None则加载所有支持的文档
            
        Returns:
            LlamaDocument列表
        """
        documents = []
        
        if file_paths is None:
            # 扫描知识库目录下的所有文档
            file_paths = []
            for ext in ['*.docx', '*.pdf', '*.txt']:
                file_paths.extend(self.knowledge_base_dir.glob(ext))
        
        for file_path in file_paths:
            try:
                file_path = Path(file_path)
                content = self._extract_content(file_path)
                
                if content:
                    # 创建LlamaDocument实例
                    doc = LlamaDocument(
                        text=content,
                        metadata={
                            "file_name": file_path.name,
                            "file_path": str(file_path),
                            "file_type": file_path.suffix,
                            "subject": self._extract_subject(file_path.name),
                            "grade": self._extract_grade(file_path.name),
                        }
                    )
                    documents.append(doc)
                    logger.info(f"成功加载文档: {file_path.name}")
                    
            except Exception as e:
                logger.error(f"加载文档失败 {file_path}: {e}")
                continue
        
        logger.info(f"总共加载了 {len(documents)} 个教案文档")
        return documents
    
    def _extract_content(self, file_path: Path) -> str:
        """
        从文件中提取文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取的文本内容
        """
        content = ""
        
        try:
            if file_path.suffix.lower() == '.docx':
                # 处理Word文档
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_path.suffix.lower() == '.pdf':
                # 处理PDF文档
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = '\n'.join([page.extract_text() for page in pdf_reader.pages])
                    
            elif file_path.suffix.lower() == '.txt':
                # 处理文本文件
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    
        except Exception as e:
            logger.error(f"提取文件内容失败 {file_path}: {e}")
            
        return content.strip()
    
    def _extract_subject(self, filename: str) -> str:
        """从文件名提取学科信息"""
        subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
        for subject in subjects:
            if subject in filename:
                return subject
        return "未知学科"
    
    def _extract_grade(self, filename: str) -> str:
        """从文件名提取年级信息"""
        grades = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级',
                 '七年级', '八年级', '九年级', '高一', '高二', '高三']
        for grade in grades:
            if grade in filename:
                return grade
        return "未知年级"
    
    def build_index(self, documents: List[LlamaDocument] = None) -> VectorStoreIndex:
        """
        构建向量索引
        
        Args:
            documents: 要索引的文档列表，如果为None则加载所有文档
            
        Returns:
            向量索引实例
        """
        if documents is None:
            documents = self.load_documents()
        
        if not documents:
            raise ValueError("没有找到可用的文档来构建索引")
        
        try:
            # 创建文本分割器
            text_splitter = SentenceSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap
            )
            
            # 构建索引
            self.index = VectorStoreIndex.from_documents(
                documents,
                storage_context=self.storage_context,
                transformations=[text_splitter]
            )
            
            logger.info(f"成功构建向量索引，包含 {len(documents)} 个文档")
            return self.index
            
        except Exception as e:
            logger.error(f"构建向量索引失败: {e}")
            raise
    
    def load_existing_index(self) -> Optional[VectorStoreIndex]:
        """
        加载已存在的向量索引
        
        Returns:
            向量索引实例或None
        """
        try:
            self.index = VectorStoreIndex.from_vector_store(
                vector_store=self.vector_store,
                storage_context=self.storage_context
            )
            logger.info("成功加载已存在的向量索引")
            return self.index
            
        except Exception as e:
            logger.warning(f"加载已存在索引失败: {e}")
            return None
    
    def search_similar_lessons(self, query: str, top_k: int = None) -> List[Dict[str, Any]]:
        """
        检索相似教案
        
        Args:
            query: 查询文本
            top_k: 返回的最相似文档数量
            
        Returns:
            相似教案列表
        """
        if top_k is None:
            top_k = settings.similarity_top_k
        
        if not hasattr(self, 'index'):
            # 尝试加载已存在的索引
            if self.load_existing_index() is None:
                # 如果没有索引，构建新的
                self.build_index()
        
        try:
            # 创建查询引擎
            query_engine = self.index.as_query_engine(
                similarity_top_k=top_k,
                response_mode="no_text"  # 只返回节点，不生成回答
            )
            
            # 执行查询
            response = query_engine.query(query)
            
            # 处理结果
            results = []
            for node in response.source_nodes:
                results.append({
                    "content": node.text,
                    "score": node.score,
                    "metadata": node.metadata,
                    "file_name": node.metadata.get("file_name", "未知文件"),
                    "subject": node.metadata.get("subject", "未知学科"),
                    "grade": node.metadata.get("grade", "未知年级")
                })
            
            logger.info(f"检索到 {len(results)} 个相似教案")
            return results
            
        except Exception as e:
            logger.error(f"检索相似教案失败: {e}")
            return []
    
    def add_lesson_plan(self, file_path: str, metadata: Dict[str, Any] = None) -> bool:
        """
        添加新的教案到知识库
        
        Args:
            file_path: 教案文件路径
            metadata: 额外的元数据
            
        Returns:
            是否添加成功
        """
        try:
            # 加载单个文档
            documents = self.load_documents([file_path])
            
            if not documents:
                return False
            
            # 更新元数据
            if metadata:
                documents[0].metadata.update(metadata)
            
            # 如果索引不存在，创建新的
            if not hasattr(self, 'index'):
                self.index = self.build_index(documents)
            else:
                # 将新文档添加到现有索引
                self.index.insert(documents[0])
            
            logger.info(f"成功添加教案: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"添加教案失败: {e}")
            return False
    
    def get_knowledge_base_stats(self) -> Dict[str, Any]:
        """
        获取知识库统计信息
        
        Returns:
            统计信息字典
        """
        try:
            # 获取collection信息
            collection_count = self.chroma_collection.count()
            
            # 扫描文件系统中的文档
            doc_files = []
            for ext in ['*.docx', '*.pdf', '*.txt']:
                doc_files.extend(self.knowledge_base_dir.glob(ext))
            
            # 按学科分类统计
            subject_stats = {}
            grade_stats = {}
            
            for file_path in doc_files:
                subject = self._extract_subject(file_path.name)
                grade = self._extract_grade(file_path.name)
                
                subject_stats[subject] = subject_stats.get(subject, 0) + 1
                grade_stats[grade] = grade_stats.get(grade, 0) + 1
            
            return {
                "total_documents": len(doc_files),
                "indexed_chunks": collection_count,
                "subjects": subject_stats,
                "grades": grade_stats,
                "knowledge_base_dir": str(self.knowledge_base_dir),
                "vector_store_dir": str(self.chroma_persist_dir)
            }
            
        except Exception as e:
            logger.error(f"获取知识库统计信息失败: {e}")
            return {}

# 创建全局知识库实例
knowledge_base = LessonPlanKnowledgeBase()